from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_block(document, block_data):
    '''Adds a block of text to a docx document.'''

    paragraph = document.add_paragraph()

    # Accessing JSON data safely
    text = block_data.get('text', '')  # Default to empty string if 'text' is missing

    # Applying styles
    run = paragraph.add_run(text)

    font_family = block_data.get('font-family', 'Calibri')  # Default font
    font_size = block_data.get('font-size', '11pt')  # Default font size

    try:
        font_size_pt = int(font_size)
    except (ValueError, TypeError):
        font_size_pt = 11

    run.font.name = font_family
    run.font.size = Pt(font_size_pt)
    run.font.bold = block_data.get('font-weight', 'normal') == 'bold'
    

    #Color Handling
    color_hex = block_data.get('color', None)
    if color_hex:
        try:
            r, g, b = tuple(int(color_hex.lstrip('#')[i:i+2], 16) for i in (0, 2, 4))
            run.font.color.rgb = RGBColor(r,g,b)
        except (ValueError, TypeError):
            pass # Ignore invalid color hex


    alignment = block_data.get('alignment', 'left')
    if alignment == 'center':
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER


def create_page(document, page_data):
    '''Creates a docx document from a single JSON object.'''

    for block_data in page_data:
        add_block(document, block_data)


def create_document(json_data):
    '''Creates a docx document from a list of JSON objects.'''

    document = Document()
    for page_data in json_data:
        create_page(document, page_data)
        document.add_page_break()

    document.save('output.docx')